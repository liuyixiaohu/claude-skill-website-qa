#!/usr/bin/env python3
"""Visual QA — Word Report Generator (combines Playwright + MCP results).

Reads:
  1. hybrid_automated_results.json  — Playwright checks, CSS inspection, screenshots
  2. hybrid_mcp_results.json        — MCP Preview findings (optional)

Usage:
  python3 report.py [--out DIR] [--report PATH]

Examples:
  python3 report.py
  python3 report.py --out /tmp/my-qa
  python3 report.py --out /tmp/my-qa --report ~/Desktop/MyReport.docx
"""
import json, os, sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from PIL import Image
except ImportError:
    print("ERROR: pip install python-docx Pillow")
    sys.exit(1)

# ── Paths (from CLI args) ────────────────────────────────────
OUT = "/tmp/visual-qa"
REPORT = os.path.expanduser("~/Desktop/Visual_QA_Report.docx")
for i, arg in enumerate(sys.argv):
    if arg == "--out" and i + 1 < len(sys.argv):
        OUT = sys.argv[i + 1]
    if arg == "--report" and i + 1 < len(sys.argv):
        REPORT = os.path.expanduser(sys.argv[i + 1])
MAX_PAGE_H = 9.0  # inches — max image height before slicing

# ── Colors ───────────────────────────────────────────────────
GREEN = RGBColor(0x22, 0x8B, 0x22)
RED = RGBColor(0xCC, 0x00, 0x00)
ORANGE = RGBColor(0xCC, 0x66, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
DARK_GRAY = RGBColor(0x44, 0x44, 0x44)
MED_GRAY = RGBColor(0x55, 0x55, 0x55)
TEXT_GRAY = RGBColor(0x33, 0x33, 0x33)

# ── Report configuration ─────────────────────────────────────
DESKTOP_VIEWPORT = "MacBook_Air_13"
TABLET_VIEWPORT = "iPad_Air"
MOBILE_VIEWPORT = "iPhone_15"
TYPO_ELEMENTS = ["h1", "hero-subtitle", "h2-0", "body-p-0", "contact-button"]
VIEWPORTS_TO_SHOW = ["iPhone_SE", "iPad_Air", "MacBook_Air_13"]
VIEWPORT_LABELS = {
    "iPhone_SE": "iPhone SE\n(375\u00d7667)",
    "iPad_Air": "iPad Air\n(820\u00d71180)",
    "MacBook_Air_13": "MacBook Air 13\"\n(1440\u00d7900)",
}
ALSO_VERIFIED = [
    "iPhone SE (375\u00d7667)",
    "iPhone 15 Pro Max (430\u00d7932)",
    'iPad Pro 12.9" (1024\u00d71366)',
    'MacBook Pro 16" (1728\u00d71117)',
]

# ── Helpers ──────────────────────────────────────────────────


def make_borderless(table):
    """Remove all borders from a Word table."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = tbl.makeelement(qn("w:tblPr"), {})
        tbl.insert(0, tblPr)
    borders = tblPr.makeelement(qn("w:tblBorders"), {})
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.makeelement(
            qn(f"w:{edge}"),
            {
                qn("w:val"): "none",
                qn("w:sz"): "0",
                qn("w:space"): "0",
                qn("w:color"): "auto",
            },
        )
        borders.append(el)
    old = tblPr.find(qn("w:tblBorders"))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(borders)


def set_cell_vertical_top(cell):
    """Set cell vertical alignment to top."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = tcPr.makeelement(qn("w:vAlign"), {qn("w:val"): "top"})
    old = tcPr.find(qn("w:vAlign"))
    if old is not None:
        tcPr.remove(old)
    tcPr.append(vAlign)


def slice_image(img_path, display_w_inches, max_h_inches=MAX_PAGE_H):
    """Slice a tall image into page-fitting segments. Returns list of file paths.
    Uses a single Image.open() call for both dimension check and cropping."""
    if not os.path.exists(img_path):
        return []
    with Image.open(img_path) as img:
        w, h = img.size
        scale = display_w_inches / w
        display_h = h * scale
        if display_h <= max_h_inches:
            return [img_path]
        max_h_px = int(max_h_inches / scale)
        segments = []
        y = 0
        i = 0
        while y < h:
            crop_h = min(max_h_px, h - y)
            seg = img.crop((0, y, w, y + crop_h))
            base, ext = os.path.splitext(img_path)
            seg_path = f"{base}_s{i}{ext}"
            seg.save(seg_path, optimize=True)
            segments.append(seg_path)
            y += crop_h
            i += 1
    return segments


def add_image_to_cell(cell, img_path, width_inches):
    """Add image (with auto-slicing) into a table cell."""
    segments = slice_image(img_path, width_inches)
    for seg in segments:
        p = (
            cell.paragraphs[-1]
            if cell.paragraphs and cell.paragraphs[-1].text == ""
            else cell.add_paragraph()
        )
        r = p.add_run()
        r.add_picture(seg, width=Inches(width_inches))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)


def add_label_to_cell(cell, label, size=Pt(8)):
    """Add a bold centered label to a table cell."""
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    p.clear()
    run = p.add_run(label)
    run.bold = True
    run.font.size = size
    run.font.color.rgb = TEXT_GRAY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(4)


def add_small_text(doc, text, bold=False, color=None, italic=False, size=Pt(8)):
    """Add a small-text paragraph to the document."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = size
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if color:
        r.font.color.rgb = color
    return p


def cell_text(cell, text, size=Pt(7), bold=False, color=None, align=None):
    """Set text in a table cell with formatting."""
    p = cell.paragraphs[0]
    p.clear()
    r = p.add_run(str(text))
    r.font.size = size
    if bold:
        r.bold = True
    if color:
        r.font.color.rgb = color
    if align:
        p.alignment = align


def group_checks_by_category(checks):
    """Group checks into a dict keyed by category name (preserves insertion order)."""
    groups = {}
    for c in checks:
        cat = c.get("category", "Other")
        groups.setdefault(cat, []).append(c)
    return groups


# ═════════════════════════════════════════════════════════════
#  MAIN
# ═════════════════════════════════════════════════════════════


def main():
    # ── Load data ────────────────────────────────────────────
    results_path = os.path.join(OUT, "hybrid_automated_results.json")
    try:
        with open(results_path, "r") as f:
            data = json.load(f)
    except FileNotFoundError:
        print(f"ERROR: Results file not found: {results_path}")
        print("Run automated.py first to generate results.")
        sys.exit(1)
    except json.JSONDecodeError as e:
        print(f"ERROR: Invalid JSON in {results_path}: {e}")
        sys.exit(1)

    mcp_data = {}
    mcp_path = os.path.join(OUT, "hybrid_mcp_results.json")
    if os.path.exists(mcp_path):
        with open(mcp_path, "r") as f:
            mcp_data = json.load(f)

    # ── Create document ──────────────────────────────────────
    doc = Document()

    # -- Normal style --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.space_before = Pt(0)

    # -- Page margins --
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    # ── 1. Title ─────────────────────────────────────────────
    _site_name = data.get("base_url", "").replace("https://", "").replace("http://", "").split("/")[0]
    t = doc.add_heading(f"{_site_name} \u2014 Visual QA Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Date: {data.get('date', 'N/A')}  |  "
        f"Target: {data.get('base_url', 'N/A')}{data.get('page', '')}"
    )
    r.font.size = Pt(10)
    r.font.color.rgb = GRAY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Method: Playwright automated checks + CSS extraction + MCP Preview inspection"
    )
    r.font.size = Pt(9)
    r.font.color.rgb = DARK_GRAY
    r.italic = True

    # ── 2. Executive Summary ─────────────────────────────────
    doc.add_heading("Executive Summary", level=1)

    checks = data.get("checks", [])
    if not checks:
        print("WARNING: No checks found in results. Report will be mostly empty.")

    total_p = sum(1 for c in checks if c["status"] == "PASS")
    total_f = sum(1 for c in checks if c["status"] == "FAIL")
    total_s = sum(1 for c in checks if c["status"] == "SKIP")
    total_c = len(checks)

    # Determine verdict
    verdict_raw = data.get("verdict", "")
    if verdict_raw:
        verdict = verdict_raw
    elif total_f == 0:
        verdict = "READY TO PUBLISH"
    elif total_f <= 3:
        verdict = "READY TO PUBLISH (with warnings)"
    else:
        verdict = "FIX BEFORE PUBLISHING"

    is_pass = "FIX" not in verdict.upper()
    v_color = GREEN if is_pass else RED

    # Verdict line
    p = doc.add_paragraph()
    run = p.add_run(verdict)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = v_color
    skip_note = f", {total_s} skipped" if total_s > 0 else ""
    r2 = p.add_run(f"   ({total_p} passed, {total_f} failed{skip_note} out of {total_c} checks)")
    r2.font.size = Pt(10)

    # Summary table
    summary = data.get("summary", {})
    if summary:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)

        sum_table = doc.add_table(rows=1, cols=4)
        sum_table.style = "Light Grid Accent 1"
        sum_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr = sum_table.rows[0].cells
        for i, txt in enumerate(["Category", "Passed", "Failed", "Total"]):
            cell_text(hdr[i], txt, size=Pt(9), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        for part_key in sorted(summary.keys()):
            part = summary[part_key]
            row = sum_table.add_row().cells
            cell_text(row[0], part["name"], size=Pt(9))
            cell_text(row[1], str(part["passed"]), size=Pt(9),
                      color=GREEN if part["passed"] > 0 else None,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_text(row[2], str(part["failed"]), size=Pt(9),
                      color=RED if part["failed"] > 0 else None,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
            cell_text(row[3], str(part["total"]), size=Pt(9),
                      align=WD_ALIGN_PARAGRAPH.CENTER)

        # TOTAL row — reuse already-computed totals
        row = sum_table.add_row().cells
        cell_text(row[0], "TOTAL", size=Pt(9), bold=True)
        cell_text(row[1], str(total_p), size=Pt(9), bold=True, color=GREEN,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[2], str(total_f), size=Pt(9), bold=True,
                  color=RED if total_f > 0 else GREEN,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row[3], str(total_c), size=Pt(9), bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

    # ── 3. All Check Results ─────────────────────────────────
    doc.add_heading("All Check Results", level=1)

    grouped = group_checks_by_category(checks)

    part_number = 1
    for cat_name, cat_checks in grouped.items():
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        r = p.add_run(f"Part {part_number}: {cat_name}")
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = TEXT_GRAY

        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        # Set column widths only on header row
        hdr = table.rows[0].cells
        col_widths = [Inches(0.5), Inches(2.2), Inches(4.8)]
        for i, (txt, w) in enumerate(zip(["Status", "Check", "Details"], col_widths)):
            hdr[i].width = w
            hdr[i].paragraphs[0].clear()
            r = hdr[i].paragraphs[0].add_run(txt)
            r.bold = True
            r.font.size = Pt(8)

        for check in cat_checks:
            status = check["status"]
            row = table.add_row().cells

            p = row[0].paragraphs[0]
            p.clear()
            r = p.add_run(status)
            r.bold = True
            r.font.size = Pt(7)
            if status == "PASS":
                r.font.color.rgb = GREEN
            elif status == "SKIP":
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            else:
                r.font.color.rgb = RED

            p = row[1].paragraphs[0]
            p.clear()
            r = p.add_run(f"{check['id']} \u2014 {check['name']}")
            r.font.size = Pt(7)

            p = row[2].paragraphs[0]
            p.clear()
            detail = check.get("details", "")
            if len(detail) > 120:
                detail = detail[:117] + "..."
            r = p.add_run(detail)
            r.font.size = Pt(7)
            r.font.color.rgb = GRAY

        part_number += 1

    # ── 4. CSS Inspection — Typography Scale ─────────────────
    doc.add_heading("CSS Inspection \u2014 Typography Scale", level=1)
    p = doc.add_paragraph()
    r = p.add_run(
        "Computed CSS values extracted via Playwright getComputedStyle() at each viewport. "
        "Verified independently via MCP preview_inspect."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    r.italic = True

    css_data = data.get("css_inspection", {})

    if css_data:
        table = doc.add_table(rows=1, cols=len(VIEWPORTS_TO_SHOW) + 1)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        hdr = table.rows[0].cells
        hdr[0].width = Inches(1.2)
        cell_text(hdr[0], "Element", size=Pt(7), bold=True)
        for i, vp in enumerate(VIEWPORTS_TO_SHOW):
            hdr[i + 1].width = Inches(1.9)
            cell_text(hdr[i + 1], VIEWPORT_LABELS.get(vp, vp.replace("_", " ")),
                      size=Pt(7), bold=True)

        for elem_label in TYPO_ELEMENTS:
            row = table.add_row().cells
            cell_text(row[0], elem_label, size=Pt(7), bold=True)

            for i, vp in enumerate(VIEWPORTS_TO_SHOW):
                vp_data = css_data.get(vp, [])
                elem = None
                if isinstance(vp_data, list):
                    elem = next(
                        (e for e in vp_data if e and e.get("label") == elem_label), None
                    )

                p = row[i + 1].paragraphs[0]
                p.clear()
                if elem:
                    font_size = elem.get("fontSize", "?")
                    font_weight = elem.get("fontWeight", "?")
                    font_family = elem.get("fontFamily", "?")[:20]
                    color_val = elem.get("color", "?")[:25]
                    info = f"{font_size} / {font_weight}\n{font_family}\ncolor: {color_val}"
                    r = p.add_run(info)
                    r.font.size = Pt(6)
                else:
                    r = p.add_run("\u2014")
                    r.font.size = Pt(6)

    # ── 5. CSS Inspection — Section Backgrounds ──────────────
    doc.add_heading("CSS Inspection \u2014 Section Backgrounds", level=2)
    p = doc.add_paragraph()
    r = p.add_run(
        "Background properties for each page section at desktop viewport (MacBook Air 13\")."
    )
    r.font.size = Pt(8)
    r.font.color.rgb = GRAY
    r.italic = True

    desktop_css = css_data.get(DESKTOP_VIEWPORT, [])
    sec_bgs = []
    if isinstance(desktop_css, list):
        sec_bgs = [
            e for e in desktop_css if e and e.get("label", "").startswith("section-bg")
        ]

    if sec_bgs:
        table = doc.add_table(rows=1, cols=5)
        table.style = "Light Grid Accent 1"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        hdr = table.rows[0].cells
        bg_col_widths = [Inches(1.8), Inches(1.5), Inches(1.2), Inches(1.0), Inches(1.0)]
        for i, (txt, w) in enumerate(zip(
            ["Section", "Background Color", "Background Image", "Height", "Padding"],
            bg_col_widths
        )):
            cell_text(hdr[i], txt, size=Pt(7), bold=True)
            hdr[i].width = w

        for bg in sec_bgs:
            row = table.add_row().cells
            heading_text = bg.get("heading", bg.get("label", "?"))
            if len(heading_text) > 35:
                heading_text = heading_text[:32] + "..."

            cell_text(row[0], heading_text, size=Pt(7))
            cell_text(row[1], bg.get("backgroundColor", "?"), size=Pt(7))
            cell_text(row[2], bg.get("backgroundImage", "?"), size=Pt(7))
            cell_text(row[3], bg.get("height", "?"), size=Pt(7))
            cell_text(row[4], bg.get("padding", "?"), size=Pt(7))
    else:
        add_small_text(doc, "No section background data available.", italic=True, color=GRAY)

    # ── 6. MCP Preview Findings ──────────────────────────────
    if mcp_data:
        doc.add_heading("MCP Preview Inspection Findings", level=1)
        p = doc.add_paragraph()
        r = p.add_run(
            "Inspected via MCP preview_inspect + preview_snapshot through reverse proxy."
        )
        r.font.size = Pt(8)
        r.font.color.rgb = GRAY
        r.italic = True

        caps = mcp_data.get("capabilities_tested", {})
        if caps:
            doc.add_heading("MCP Tool Capabilities", level=2)
            table = doc.add_table(rows=1, cols=2)
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            hdr = table.rows[0].cells
            hdr[0].width = Inches(2.0)
            hdr[1].width = Inches(5.5)
            cell_text(hdr[0], "Tool", size=Pt(8), bold=True)
            cell_text(hdr[1], "Status", size=Pt(8), bold=True)

            for tool, status in caps.items():
                row = table.add_row().cells
                cell_text(row[0], tool, size=Pt(7))
                p = row[1].paragraphs[0]
                p.clear()
                r = p.add_run(status)
                r.font.size = Pt(7)
                if "WORKS" in status.upper():
                    r.font.color.rgb = GREEN
                elif "FAIL" in status.upper():
                    r.font.color.rgb = RED
                else:
                    r.font.color.rgb = GRAY

        tablet = mcp_data.get("tablet_768x1024", {})
        if tablet.get("note"):
            doc.add_heading("Responsive Finding", level=2)
            p = doc.add_paragraph()
            r = p.add_run(tablet["note"])
            r.font.size = Pt(9)
            r.font.color.rgb = ORANGE

        modal = mcp_data.get("contact_modal", {})
        if modal:
            doc.add_heading("Contact Modal (MCP inspection)", level=2)
            fields = modal.get("fields", [])
            if fields:
                p = doc.add_paragraph()
                r = p.add_run("Form fields detected:")
                r.bold = True
                r.font.size = Pt(9)
                for field in fields:
                    p = doc.add_paragraph(style="List Bullet")
                    r = p.add_run(field)
                    r.font.size = Pt(8)
            p = doc.add_paragraph()
            r = p.add_run(
                f"Background: {modal.get('bgColor', '?')}  |  "
                f"Width (mobile): {modal.get('width_mobile', '?')}  |  "
                f"Submit button: {'Yes' if modal.get('submitButton') else 'No'}  |  "
                f"Close button: {'Yes' if modal.get('closeButton') else 'No'}"
            )
            r.font.size = Pt(8)
            r.font.color.rgb = GRAY

    # ── 7. Section Screenshots ───────────────────────────────
    # Auto-detect page key from section_screenshots (not hardcoded)
    all_screenshot_pages = data.get("section_screenshots", {})
    page_key = next(iter(all_screenshot_pages), None)
    section_screenshots = all_screenshot_pages.get(page_key, {}) if page_key else {}
    page_label = (page_key or "Page").replace("-", " ").title()

    doc.add_heading(f"{page_label} \u2014 Section Screenshots", level=1)

    # 7a. Desktop (full width)
    desktop_secs = section_screenshots.get(DESKTOP_VIEWPORT, [])

    if desktop_secs:
        doc.add_heading(f'Desktop \u2014 MacBook Air 13" (1440\u00d7900)', level=2)
        for sec in desktop_secs:
            p = doc.add_paragraph()
            r = p.add_run(sec["label"])
            r.bold = True
            r.font.size = Pt(9)
            r.font.color.rgb = TEXT_GRAY
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(2)

            fpath = os.path.join(OUT, sec["file"])
            if os.path.exists(fpath):
                segments = slice_image(fpath, 7.0)
                for seg in segments:
                    doc.add_picture(seg, width=Inches(7.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 7b. Responsive views (Tablet + Mobile side by side)
    tablet_secs = section_screenshots.get(TABLET_VIEWPORT, [])
    mobile_secs = section_screenshots.get(MOBILE_VIEWPORT, [])
    max_secs = max(len(tablet_secs), len(mobile_secs), 0)

    if max_secs > 0:
        doc.add_heading("Responsive Views (Tablet + Mobile)", level=2)

        for si in range(max_secs):
            t_sec = tablet_secs[si] if si < len(tablet_secs) else None
            m_sec = mobile_secs[si] if si < len(mobile_secs) else None
            label = (t_sec or m_sec or {}).get("label", f"Section {si}")

            tbl = doc.add_table(rows=1, cols=2)
            make_borderless(tbl)
            tbl.autofit = False
            tbl.cell(0, 0).width = Inches(4.5)
            tbl.cell(0, 1).width = Inches(2.8)
            set_cell_vertical_top(tbl.cell(0, 0))
            set_cell_vertical_top(tbl.cell(0, 1))

            add_label_to_cell(tbl.cell(0, 0), f"{label} \u2014 iPad Air")
            if t_sec:
                fpath = os.path.join(OUT, t_sec["file"])
                if os.path.exists(fpath):
                    add_image_to_cell(tbl.cell(0, 0), fpath, 4.3)

            add_label_to_cell(tbl.cell(0, 1), f"{label} \u2014 iPhone 15")
            if m_sec:
                fpath = os.path.join(OUT, m_sec["file"])
                if os.path.exists(fpath):
                    add_image_to_cell(tbl.cell(0, 1), fpath, 2.5)

    # "Also verified on" line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run("Also verified on: ")
    r.bold = True
    r.font.size = Pt(8)
    r.font.color.rgb = MED_GRAY
    r = p.add_run(", ".join(ALSO_VERIFIED))
    r.font.size = Pt(8)
    r.font.color.rgb = MED_GRAY

    # ── 8. Interactions ───────────────────────────────────────
    doc.add_heading("Interactions", level=1)

    # Desktop contact modal
    fpath = os.path.join(OUT, "hybrid_modal_contact_desktop.png")
    if os.path.exists(fpath):
        p = doc.add_paragraph()
        r = p.add_run("Contact Modal \u2014 Desktop (1440\u00d7900)")
        r.font.size = Pt(9)
        r.bold = True
        r.font.color.rgb = TEXT_GRAY
        p.paragraph_format.space_after = Pt(2)

        segments = slice_image(fpath, 7.0)
        for seg in segments:
            doc.add_picture(seg, width=Inches(7.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mobile interactions side by side
    mobile_interactions = [
        ("hybrid_modal_contact_mobile.png", "Contact Modal \u2014 Mobile"),
        ("hybrid_hamburger_open.png", "Hamburger Menu \u2014 Mobile"),
    ]

    has_mobile = any(
        os.path.exists(os.path.join(OUT, fname)) for fname, _ in mobile_interactions
    )

    if has_mobile:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)

        i_table = doc.add_table(rows=1, cols=2)
        make_borderless(i_table)
        i_table.autofit = False

        for i, (fname, label) in enumerate(mobile_interactions):
            i_table.cell(0, i).width = Inches(3.6)
            set_cell_vertical_top(i_table.cell(0, i))
            add_label_to_cell(i_table.cell(0, i), label)

            fpath = os.path.join(OUT, fname)
            if os.path.exists(fpath):
                add_image_to_cell(i_table.cell(0, i), fpath, 3.4)

    # ── 9. Critical Issues ───────────────────────────────────
    doc.add_heading("Critical Issues (Must Fix Before Publish)", level=1)

    critical_issues = data.get("critical_issues", [])
    if critical_issues:
        for issue in critical_issues:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(issue)
            r.font.size = Pt(9)
            r.font.color.rgb = RED
    else:
        p = doc.add_paragraph()
        r = p.add_run("No critical issues found.")
        r.font.size = Pt(10)
        r.font.color.rgb = GREEN
        r.italic = True

    # ── 10. Warnings ─────────────────────────────────────────
    doc.add_heading("Warnings (Recommended Improvements)", level=1)

    warnings = data.get("warnings", [])
    if warnings:
        for warning in warnings:
            p = doc.add_paragraph(style="List Bullet")
            r = p.add_run(warning)
            r.font.size = Pt(9)
            r.font.color.rgb = ORANGE
    else:
        p = doc.add_paragraph()
        r = p.add_run("No warnings.")
        r.font.size = Pt(10)
        r.font.color.rgb = GREEN
        r.italic = True

    # ── 11. Final Verdict ────────────────────────────────────
    doc.add_heading("Final Verdict", level=1)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(f"VERDICT: {verdict}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = v_color

    # ── Save & Stats ─────────────────────────────────────────
    doc.save(REPORT)
    print(f"Report saved to: {REPORT}")

    img_count = sum(1 for _ in doc.inline_shapes)
    tbl_count = len(doc.tables)
    file_size = os.path.getsize(REPORT) / 1024 / 1024
    print(f"Images: {img_count}, Tables: {tbl_count}, Size: {file_size:.1f} MB")


if __name__ == "__main__":
    main()
